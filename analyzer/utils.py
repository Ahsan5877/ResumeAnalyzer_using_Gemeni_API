import docx
import fitz
import PyPDF2
from docx import Document

def extract_text_from_file(file):
    if file.name.endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        return " ".join([page.extract_text() for page in reader.pages])
    elif file.name.endswith('.docx'):
        doc = Document(file)
        return " ".join([para.text for para in doc.paragraphs])
    return ""

def extract_pdf_text(path):
    text = ""
    with fitz.open(path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_docx_text(path):
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])