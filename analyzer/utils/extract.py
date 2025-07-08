import docx
import fitz
import PyPDF2
from docx import Document

def extract_text_from_file(filepath):
    """Extract text from PDF or DOCX files"""
    try:
        if filepath.lower().endswith('.pdf'):
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return " ".join([page.extract_text() for page in reader.pages])
                
        elif filepath.lower().endswith('.docx'):
            doc = Document(filepath)
            return " ".join([para.text for para in doc.paragraphs])
            
        return "Unsupported file format"
    except Exception as e:
        print(f"Extraction error: {str(e)}")
        return "Failed to extract text"

def extract_pdf_text(path):
    text = ""
    with fitz.open(path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_docx_text(path):
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])